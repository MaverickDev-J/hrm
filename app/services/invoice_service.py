import os
from datetime import datetime, date
from uuid import UUID
from typing import List, Dict, Any, Optional

from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.invoice import Invoice
from app.models.client import Client
from app.models.company import Company
from app.models.candidate import Candidate
from app.services.client_service import get_client_column_config
from app.schemas.invoice import ManualTotals

# Paths
STATIC_DIR = "static"
INVOICE_DIR = os.path.join(STATIC_DIR, "invoices")
os.makedirs(INVOICE_DIR, exist_ok=True)

# --- Helper Functions ---
def add_border_to_table(table, border_size='4'):
    """Add borders to table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_background(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading)

def set_cell_vertical_alignment(cell, align="center"):
    """Set vertical alignment for cell"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def set_repeat_table_header(row):
    """Set table row to repeat as header on new pages"""
    tr = row._element
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

class InvoiceGenerator:
    def __init__(self, db: Session):
        self.db = db

    def prepare_invoice_data(
        self, 
        company_id: UUID, 
        client_id: UUID, 
        candidate_ids: List[UUID], 
        manual_totals: ManualTotals,
        invoice_number: str,
        invoice_date: date
    ) -> Dict[str, Any]:
        """
        Aggregates all necessary data for invoice generation.
        """
        # 1. Company Data
        company = self.db.query(Company).filter(Company.id == company_id).first()
        if not company:
            raise ValueError("Company not found")

        # 2. Client Data
        client = self.db.query(Client).filter(Client.id == client_id).first()
        if not client:
            raise ValueError("Client not found")

        # 3. Candidates Data
        candidates = self.db.query(Candidate).filter(Candidate.id.in_(candidate_ids)).all()

        # 4. Column Config
        config = get_client_column_config(self.db, client_id)
        if config and config.column_definitions:
            raw_columns = config.column_definitions.get("columns", [])
            columns = []
            for col in raw_columns:
                # Filter out serial no as it is auto-generated
                fname = col.get("field_name", "").lower()
                if fname in ["sr_no", "serial_no", "s_no", "s.no"]:
                    continue

                # Normalize column data
                col_def = col.copy()
                # Map column_width to width if present, else default
                w = col.get("width") or col.get("column_width") or 1.0
                col_def["width"] = float(w)
                columns.append(col_def)
        else:
            columns = []
        
        # If no config, default to basic columns
        if not columns:
            columns = [
                {"field_name": "candidate_name", "display_label": "Candidate Name", "width": 2.0},
                {"field_name": "amount", "display_label": "Amount", "width": 1.0}
            ]

        # 5. Structure Line Items
        line_items = []
        for index, cand in enumerate(candidates, start=1):
            item = {"serial_no": index}
            data = cand.candidate_data
            if not data: 
                data = {}
            
            # Map columns
            for col in columns:
                fname = col.get("field_name")
                val = data.get(fname, "")
                item[fname] = val
            
            # Ensure amount is accessible
            item["amount"] = data.get("amount", 0)
            line_items.append(item)

        # 6. Return Structured Dict
        return {
            "invoice_number": invoice_number,
            "invoice_date": invoice_date.strftime("%d-%b-%Y"),
            "company": {
                "name": company.name,
                "tagline": getattr(company, 'tagline', '') or "",
                "address_line1": getattr(company, 'address_line1', '') or "",
                "city": getattr(company, 'city', '') or "",
                "state": getattr(company, 'state', '') or "",
                "pincode": getattr(company, 'pincode', '') or "",
                "pan": getattr(company, 'pan', '') or getattr(company, 'pan_number', '') or "",
                "banner_url": getattr(company, 'banner_image_url', None), # URL can be None
                "stamp_url": getattr(company, 'stamp_url', None),
                "signature_url": getattr(company, 'signature_url', None),
                "bank_name": getattr(company, 'bank_name', '') or "",
                "account_holder_name": getattr(company, 'account_holder_name', '') or "",
                "account_number": getattr(company, 'account_number', '') or "",
                "ifsc_code": getattr(company, 'ifsc_code', '') or "",
            },
            "client": {
                "name": client.client_name,
                "address": client.client_address,
                "address_line2": getattr(client, 'address_line2', '') or "",
                "city": getattr(client, 'city', '') or "",
                "state": getattr(client, 'state', '') or "",
                "pincode": getattr(client, 'pincode', '') or "",
                "gstin": client.gstin,
                "pan": getattr(client, 'pan', '') or getattr(client, 'pan_number', '') or "N/A"
            },
            "columns": columns,
            "line_items": line_items,
            "financials": manual_totals.model_dump()
        }

    def generate_docx(self, data: Dict[str, Any]) -> str:
        """
        Generates professional DOCX file with styling and returns relative URL.
        """
        doc = Document()
        
        # --- SET MARGINS ---
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.3)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        # --- ADD BANNER IMAGE ---
        banner_path = data["company"].get("banner_url")
        # Ensure path is relative to current working directory if it starts with /
        if banner_path and banner_path.startswith("/"):
            banner_path = "." + banner_path
            

        if banner_path and os.path.exists(banner_path):
            try:
                banner_para = doc.add_paragraph()
                banner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Standard: 7" wide x 1.2" tall
                banner_para.add_run().add_picture(banner_path, width=Inches(7), height=Inches(1.2))
                banner_para.space_after = Pt(6)
            except Exception as e:
                print(f"Could not add banner: {e}")
        
        # --- HEADER SECTION (Two Columns) ---
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.allow_autofit = False
        header_table.columns[0].width = Inches(3.5)
        header_table.columns[1].width = Inches(3.5)
        
        # LEFT: Company Details
        left_cell = header_table.rows[0].cells[0]
        left_para = left_cell.paragraphs[0]
        
        run = left_para.add_run(data['company']['name'] + '\n')
        run.bold = True
        run.font.size = Pt(14)
        
        # Add tagline, address, PAN
        if data['company'].get('tagline'):
            left_para.add_run(f"{data['company']['tagline']}\n").font.size = Pt(9)
        
        if data['company'].get('address_line1'):
            left_para.add_run(f"{data['company']['address_line1']}\n").font.size = Pt(9)
        
        city = data['company'].get('city', '')
        state = data['company'].get('state', '')
        pincode = data['company'].get('pincode', '')
        if city or state or pincode:
            left_para.add_run(f"{city}, {state} - {pincode}\n").font.size = Pt(9)
        
        left_para.add_run(f"PAN: {data['company'].get('pan', 'N/A')}").font.size = Pt(9)
        
        # RIGHT: Invoice Title and Details
        right_cell = header_table.rows[0].cells[1]
        set_cell_vertical_alignment(right_cell, "top")
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        title_run = right_para.add_run('TAX INVOICE\n\n')
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        right_para.add_run('Invoice Number: ').bold = True
        right_para.add_run(f"{data['invoice_number']}\n")
        right_para.add_run('Invoice Date: ').bold = True
        right_para.add_run(f"{data['invoice_date']}\n")
        right_para.add_run('Place of Supply: ').bold = True
        right_para.add_run(f"{data['client'].get('state', 'N/A')}\n")
        
        for run in right_para.runs[3:]:
            run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # --- BILL TO SECTION ---
        bill_to_para = doc.add_paragraph()
        bill_to_run = bill_to_para.add_run('BILL TO')
        bill_to_run.bold = True
        bill_to_run.font.size = Pt(11)
        bill_to_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Client details in styled box
        client_table = doc.add_table(rows=1, cols=1)
        client_table.autofit = False
        add_border_to_table(client_table, '8')
        
        client_cell = client_table.rows[0].cells[0]
        set_cell_background(client_cell, 'F0F0F0')  # Light gray background
        client_para = client_cell.paragraphs[0]
        
        company_run = client_para.add_run(f"{data['client']['name']}\n")
        company_run.bold = True
        company_run.font.size = Pt(12)
        
        client_para.add_run(f"{data['client']['address']}\n").font.size = Pt(10)
        
        if data['client'].get('address_line2'):
            client_para.add_run(f"{data['client']['address_line2']}\n").font.size = Pt(10)
        
        c_city = data['client'].get('city', '')
        c_state = data['client'].get('state', '')
        c_pincode = data['client'].get('pincode', '')
        if c_city or c_state or c_pincode:
            client_para.add_run(f"{c_city}, {c_state} - {c_pincode}\n").font.size = Pt(10)
        
        client_para.add_run(f"\nGSTIN: {data['client']['gstin']}  |  PAN: {data['client'].get('pan', 'N/A')}").font.size = Pt(10)
        
        doc.add_paragraph()
        
        # --- LINE ITEMS SECTION ---
        items_heading = doc.add_paragraph()
        items_run = items_heading.add_run('LINE ITEMS')
        items_run.bold = True
        items_run.font.size = Pt(11)
        items_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Create table with S.No + configured columns
        columns = data["columns"]
        num_cols = len(columns) + 1  # +1 for S.No
        
        candidates_table = doc.add_table(rows=1, cols=num_cols)
        candidates_table.autofit = False
        candidates_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_border_to_table(candidates_table)
        
        # Set column widths dynamically
        candidates_table.columns[0].width = Inches(0.4)  # S.No
        for i, col in enumerate(columns):
            width = float(col.get("width", 1.5))
            candidates_table.columns[i + 1].width = Inches(width)
        
        # HEADER ROW with dark blue background
        header_cells = candidates_table.rows[0].cells
        set_repeat_table_header(candidates_table.rows[0])  # Repeat on new pages
        
        # S.No header
        cell = header_cells[0]
        set_cell_background(cell, '2E5090')  # Dark blue
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('S.No')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        set_cell_vertical_alignment(cell, "center")
        
        # Dynamic column headers
        for i, col in enumerate(columns):
            cell = header_cells[i + 1]
            set_cell_background(cell, '2E5090')
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(col.get("display_label", col.get("field_name")))
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_vertical_alignment(cell, "center")
        
        # DATA ROWS with alternating colors
        for idx, item in enumerate(data["line_items"]):
            row_cells = candidates_table.add_row().cells
            
            # Alternating row colors (light gray)
            if idx % 2 == 0:
                for cell in row_cells:
                    set_cell_background(cell, 'F9F9F9')
            
            # S.No
            row_cells[0].text = str(item["serial_no"])
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Dynamic columns
            for i, col in enumerate(columns):
                fname = col.get("field_name")
                val = item.get(fname, "")
                
                # Format amount with rupee symbol if it's amount field
                if fname == "amount" or "amount" in fname.lower():
                    try:
                        row_cells[i + 1].text = f"₹{float(val):,.2f}"
                    except:
                        row_cells[i + 1].text = str(val)
                    row_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    row_cells[i + 1].text = str(val)
            
            # Set font size and vertical alignment for all cells
            for cell in row_cells:
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                set_cell_vertical_alignment(cell, "center")
        
        doc.add_paragraph()
        
        # --- FINANCIAL SUMMARY ---
        summary_table = doc.add_table(rows=0, cols=2)
        summary_table.autofit = False
        summary_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        add_border_to_table(summary_table)
        
        summary_table.columns[0].width = Inches(2.0)
        summary_table.columns[1].width = Inches(1.5)
        
        totals = data['financials']
        
        # Helper to add summary rows
        def add_summary_row(label, value, is_total=False):
            row = summary_table.add_row()
            label_cell = row.cells[0]
            value_cell = row.cells[1]
            
            if is_total:
                set_cell_background(label_cell, '2E5090')
                set_cell_background(value_cell, '2E5090')
            
            label_para = label_cell.paragraphs[0]
            value_para = value_cell.paragraphs[0]
            
            label_run = label_para.add_run(label)
            value_run = value_para.add_run(f"₹{float(value):,.2f}")
            
            if is_total:
                label_run.bold = True
                value_run.bold = True
                label_run.font.size = Pt(11)
                value_run.font.size = Pt(11)
                label_run.font.color.rgb = RGBColor(255, 255, 255)
                value_run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                label_run.font.size = Pt(10)
                value_run.font.size = Pt(10)
            
            label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            value_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_cell_vertical_alignment(label_cell, "center")
            set_cell_vertical_alignment(value_cell, "center")
        
        # Add rows
        add_summary_row("Subtotal", totals['subtotal'])
        
        if totals.get('cgst_amount', 0) > 0:
            cgst_rate = totals.get('cgst_rate', 0)
            add_summary_row(f"CGST @ {cgst_rate}%", totals['cgst_amount'])
        
        if totals.get('sgst_amount', 0) > 0:
            sgst_rate = totals.get('sgst_rate', 0)
            add_summary_row(f"SGST @ {sgst_rate}%", totals['sgst_amount'])
        
        if totals.get('igst_amount', 0) > 0:
            igst_rate = totals.get('igst_rate', 0)
            add_summary_row(f"IGST @ {igst_rate}%", totals['igst_amount'])
        
        add_summary_row("GRAND TOTAL", totals['grand_total'], is_total=True)
        
        doc.add_paragraph()
        
        # --- BANK DETAILS ---
        bank_heading = doc.add_paragraph()
        bank_run = bank_heading.add_run('BANK DETAILS')
        bank_run.bold = True
        bank_run.font.size = Pt(11)
        bank_run.font.color.rgb = RGBColor(0, 51, 102)
        
        bank_table = doc.add_table(rows=5, cols=2)
        bank_table.autofit = False
        add_border_to_table(bank_table)
        
        bank_table.columns[0].width = Inches(2.0)
        bank_table.columns[1].width = Inches(4.5)
        
        bank_details = [
            ('Bank Name', data['company'].get('bank_name', 'N/A')),
            ('Account Holder', data['company'].get('account_holder_name', 'N/A')),
            ('Account Number', data['company'].get('account_number', 'N/A')),
            ('IFSC Code', data['company'].get('ifsc_code', 'N/A')),
            ('PAN', data['company'].get('pan', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(bank_details):
            label_cell = bank_table.rows[i].cells[0]
            value_cell = bank_table.rows[i].cells[1]
            
            set_cell_background(label_cell, 'E8E8E8')  # Light gray for labels
            
            label_cell.text = label
            value_cell.text = value
            
            label_cell.paragraphs[0].runs[0].bold = True
            label_cell.paragraphs[0].runs[0].font.size = Pt(10)
            value_cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            set_cell_vertical_alignment(label_cell, "center")
            set_cell_vertical_alignment(value_cell, "center")
        
        doc.add_paragraph()
        
        # --- TERMS & CONDITIONS ---
        terms_para = doc.add_paragraph()
        terms_run = terms_para.add_run('Terms & Conditions: ')
        terms_run.bold = True
        terms_run.font.size = Pt(10)
        terms_para.add_run("Payment due within 30 days. Late payments subject to interest.").font.size = Pt(10)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # --- SIGNATURE AND STAMP ---
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.autofit = False
        
        # LEFT: Stamp
        stamp_cell = sig_table.rows[0].cells[0]
        stamp_para = stamp_cell.paragraphs[0]
        
        stamp_path = data['company'].get('stamp_url')
        if stamp_path and stamp_path.startswith("/"):
            stamp_path = "." + stamp_path

        if stamp_path and os.path.exists(stamp_path):
            try:
                # Standard: 1.3" x 1.3" square
                stamp_para.add_run().add_picture(stamp_path, width=Inches(1.3), height=Inches(1.3))
            except Exception as e:
                print(f"Could not add stamp: {str(e)}")
        
        # RIGHT: Signature
        sig_cell = sig_table.rows[0].cells[1]
        sig_para = sig_cell.paragraphs[0]
        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        sig_path = data['company'].get('signature_url')
        if sig_path and sig_path.startswith("/"):
            sig_path = "." + sig_path
            
        if sig_path and os.path.exists(sig_path):
            try:
                # Standard: 2.0" wide x 1.0" tall
                sig_para.add_run().add_picture(sig_path, width=Inches(2.0), height=Inches(1.0))
                sig_para.add_run('\n')
            except Exception as e:
                 print(f"Could not add signature: {str(e)}")
        
        sig_run = sig_para.add_run('Authorized Signatory')
        sig_run.bold = True
        sig_run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # --- CLOSING NOTE ---
        closing = doc.add_paragraph()
        closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
        closing_run = closing.add_run('Thank you for your business!')
        closing_run.italic = True
        closing_run.font.size = Pt(11)
        closing_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # --- SAVE DOCUMENT ---
        filename = f"{data['invoice_number']}.docx"
        file_path = os.path.join(INVOICE_DIR, filename)
        doc.save(file_path)
        
        return f"/static/invoices/{filename}"


def generate_invoice(
    db: Session,
    company_id: UUID, 
    client_id: UUID, 
    candidate_ids: List[UUID], 
    manual_totals: ManualTotals,
    invoice_number: str,
    invoice_date: date
) -> Invoice:
    generator = InvoiceGenerator(db)
    
    # 1. Aggregate
    data = generator.prepare_invoice_data(
        company_id, client_id, candidate_ids, manual_totals, invoice_number, invoice_date
    )
    
    # 2. Generate
    file_url = generator.generate_docx(data)
    
    # 3. Save Record
    db_invoice = Invoice(
        invoice_number=invoice_number,
        invoice_date=invoice_date,
        company_id=company_id,
        client_id=client_id,
        candidate_ids=[str(cid) for cid in candidate_ids],
        subtotal=manual_totals.subtotal,
        cgst_amount=manual_totals.cgst_amount,
        sgst_amount=manual_totals.sgst_amount,
        igst_amount=manual_totals.igst_amount,
        grand_total=manual_totals.grand_total,
        file_url=file_url,
        status="GENERATED"
    )
    db.add(db_invoice)
    db.commit()
    db.refresh(db_invoice)
    
    return db_invoice

def get_latest_invoice_data_by_client_id(db: Session, client_id: UUID) -> Optional[Dict[str, Any]]:
    """
    Retrieve data for the LATEST invoice generated for a specific client.
    """
    # Fetch latest invoice by date/created_at
    # Assuming invoice_date or created_at. Model has created_at? 
    # Invoice model has invoice_date (date) and usually a created_at (timestamp). 
    # Let's check model. If no created_at, use invoice_date descending and id descending.
    invoice = db.query(Invoice).filter(Invoice.client_id == client_id).order_by(Invoice.invoice_date.desc(), Invoice.id.desc()).first()
    
    if not invoice:
        return None
        
    # Reconstruct ManualTotals from stored invoice fields
    manual_totals = ManualTotals(
        subtotal=invoice.subtotal,
        cgst_amount=invoice.cgst_amount,
        sgst_amount=invoice.sgst_amount,
        igst_amount=invoice.igst_amount,
        grand_total=invoice.grand_total
    )
    
    if invoice.candidate_ids is None:
        candidate_uuids = []
    else:
        # Handle potential string vs list issue in JSONB
        if isinstance(invoice.candidate_ids, list):
             candidate_uuids = [UUID(str(cid)) for cid in invoice.candidate_ids]
        elif isinstance(invoice.candidate_ids, str):
             # Should not happen for JSONB list, but defensive coding
             candidate_uuids = [] 
        else:
             candidate_uuids = []

    generator = InvoiceGenerator(db)
    data = generator.prepare_invoice_data(
        company_id=invoice.company_id,
        client_id=invoice.client_id,
        candidate_ids=candidate_uuids,
        manual_totals=manual_totals,
        invoice_number=invoice.invoice_number,
        invoice_date=invoice.invoice_date
    )
    
    return data